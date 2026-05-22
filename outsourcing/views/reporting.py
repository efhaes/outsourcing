from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from django.conf import settings

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from xhtml2pdf import pisa

import io
import os
import calendar
from datetime import date

from outsourcing.models import (
    LaporanKegiatan, ItemKegiatan, Absensi,
    StaffSupervisor, SupervisorPerusahaan,
    Perusahaan, JenisJasa
)

from outsourcing.views.holiday_service import get_hari_libur_set


# ──────────────────────────────────────────────────
# HELPER: BUILD DATA KALENDER
# ──────────────────────────────────────────────────

def build_jadwal_kalender(item_list, bulan, tahun):
    _, days_in_month = calendar.monthrange(int(tahun), int(bulan))
    jadwal = {}
    staff_per_sub_area = {}

    for item in item_list:
        sub_area_name = item.sub_area.nama_sub_area if item.sub_area else 'Tanpa Sub Area'
        if sub_area_name not in jadwal:
            jadwal[sub_area_name] = {}
            staff_per_sub_area[sub_area_name] = set()

        key = item.pk
        if key not in jadwal[sub_area_name]:
            staff_names = (
                ', '.join([staff.nama_lengkap or staff.username for staff in item.staff.all()])
                if item.staff.exists() else '-'
            )
            jadwal[sub_area_name][key] = {
                'nama_item'    : item.nama_item,
                'task'         : item.task.nama_task if item.task else '-',
                'sub_area'     : item.sub_area.nama_sub_area if item.sub_area else '-',
                'staff'        : staff_names,
                'tanggal_aktif': set(),
            }

        for staff in item.staff.all():
            staff_per_sub_area[sub_area_name].add(staff.nama_lengkap or staff.username)

        if item.tanggal:
            jadwal[sub_area_name][key]['tanggal_aktif'].add(item.tanggal.day)

    jadwal_final = {}
    for sub_area_name, items_dict in jadwal.items():
        items_list = list(items_dict.values())
        if items_list:
            unique_staff = sorted(staff_per_sub_area[sub_area_name])
            items_list[0]['staff_unik'] = ', '.join(unique_staff)
        jadwal_final[sub_area_name] = items_list

    return jadwal_final, days_in_month


def build_absensi_kalender(absensi_list, bulan, tahun, libur_set=None):


    if libur_set is None:
        libur_set = set()

    _, days_in_month = calendar.monthrange(int(tahun), int(bulan))

    # Precompute info per hari: Minggu atau libur nasional?
    info_hari = {}
    for day in range(1, days_in_month + 1):
        tgl               = date(int(tahun), int(bulan), day)
        tgl_str           = str(tgl)
        is_minggu         = tgl.weekday() == 6       # 6 = Minggu
        is_libur_nasional = tgl_str in libur_set
        info_hari[day] = {
            'is_libur'   : is_minggu or is_libur_nasional,
            'is_minggu'  : is_minggu,
            'is_nasional': is_libur_nasional,
        }

    absensi_map = {}

    for ab in absensi_list:
        sid = ab.staff_id
        if sid not in absensi_map:
            absensi_map[sid] = {
                'nama'         : ab.staff.nama_lengkap or ab.staff.username,
                'nik'          : ab.staff.nik or '-',
                'foto'         : ab.staff.foto_profil.url if ab.staff.foto_profil else None,
                'jumlah_hadir' : 0,
                'jumlah_alpa'  : 0,
                'jumlah_izin'  : 0,
                'jumlah_cuti'  : 0,
                'jumlah_dokter': 0,
                'keterangan'   : '',
            }

        if ab.tanggal:
            status_harian = getattr(ab, 'status_harian', None)
            
            # Kalau status_harian di-set manual (Cuti, Izin, Dokter, Libur) → pakai itu
            if status_harian and status_harian.strip() and status_harian.strip() != 'P':
                status = status_harian.strip()
            else:
                # Tentukan dari status absensi
                absensi_status = ab.status  # AbsensiStatusChoices
                if absensi_status in ('masuk', 'pulang', 'terlambat', 'overtime'):
                    status = 'P'
                else:
                    # belum_absen → Alpa
                    status = 'A'

            absensi_map[sid][f'd_{ab.tanggal.day}'] = status

            if status == 'P':
                absensi_map[sid]['jumlah_hadir'] += 1
            elif status == 'A':
                absensi_map[sid]['jumlah_alpa'] += 1
            elif status == 'I':
                absensi_map[sid]['jumlah_izin'] += 1
            elif status == 'L':
                absensi_map[sid]['jumlah_cuti'] += 1
            elif status == 'DC':
                absensi_map[sid]['jumlah_dokter'] += 1

    # Build hari_list sebagai list of dict per staff
    for sid, data in absensi_map.items():
        hari_list = []
        for d in range(1, days_in_month + 1):
            status   = data.get(f'd_{d}', '')
            is_libur = info_hari[d]['is_libur']
            # Hari libur tanpa absensi → otomatis 'LB'
            if is_libur and not status:
                status = 'LB'
            hari_list.append({
                'status'  : status,
                'is_libur': is_libur,
            })
        data['hari_list'] = hari_list

    return absensi_map, days_in_month, info_hari


def build_laporan_dengan_jadwal(laporan_list, days_in_month):
    result = []
    for laporan in laporan_list:
        items_data = []
        for item in laporan.item_kegiatan.select_related('task', 'sub_area').all():
            hari_list = [''] * days_in_month
            if item.tanggal:
                idx = item.tanggal.day - 1
                if 0 <= idx < days_in_month:
                    hari_list[idx] = '✓'
            items_data.append({
                'nama_item': item.nama_item,
                'sub_area' : item.sub_area.nama_sub_area if item.sub_area else '-',
                'task'     : item.task.nama_task if item.task else '-',
                'frek'     : 'H' if item.jam_mulai else 'M',
                'hari_list': hari_list,
            })
        result.append({
            'laporan'   : laporan,
            'area_nama' : laporan.area.nama_area if laporan.area else '-',
            'supervisor': laporan.supervisor.nama_lengkap or laporan.supervisor.username,
            'items'     : items_data,
        })
    return result


# ──────────────────────────────────────────────────
# MAIN DATA GETTER
# ──────────────────────────────────────────────────

def get_data_laporan_bulanan(perusahaan_id, bulan, tahun, jenis_jasa_id):
    laporan_list = LaporanKegiatan.objects.filter(
        perusahaan_id=perusahaan_id,
        jenis_jasa_id=jenis_jasa_id,
        tanggal_laporan__year=tahun,
        tanggal_laporan__month=bulan,
    ).prefetch_related(
        'item_kegiatan__staff',
        'item_kegiatan__task',
        'item_kegiatan__sub_area',
    ).select_related('supervisor', 'area', 'jenis_jasa')

    supervisor_ids = list(
        laporan_list.values_list('supervisor_id', flat=True).distinct()
    )

    staff_list = StaffSupervisor.objects.filter(
        supervisor_id__in=supervisor_ids,
        is_active=True,
    ).select_related('staff', 'supervisor')

    staff_ids = list(staff_list.values_list('staff_id', flat=True).distinct())

    absensi_list = Absensi.objects.filter(
        staff_id__in=staff_ids,
        tanggal__year=tahun,
        tanggal__month=bulan,
    ).select_related('staff').order_by('staff__nama_lengkap', 'tanggal')

    item_list = ItemKegiatan.objects.filter(
        laporan__in=laporan_list,
    ).prefetch_related('staff').select_related(
        'task', 'sub_area', 'laporan__area'
    ).order_by('laporan__area__nama_area', 'tanggal')

 
    libur_set = get_hari_libur_set(tahun=int(tahun), bulan=int(bulan))

    jadwal_kalender, days_in_month = build_jadwal_kalender(item_list, bulan, tahun)

    absensi_kalender, _, info_hari = build_absensi_kalender(
        absensi_list, bulan, tahun, libur_set=libur_set
    )

    laporan_dengan_jadwal = build_laporan_dengan_jadwal(laporan_list, days_in_month)

    staff_subarea_map_raw = {}
    for item in item_list:
        sub_area_nama = item.sub_area.nama_sub_area if item.sub_area else '-'
        for staff in item.staff.all():
            if staff.pk not in staff_subarea_map_raw:
                staff_subarea_map_raw[staff.pk] = set()
            staff_subarea_map_raw[staff.pk].add(sub_area_nama)

    staff_subarea_map = {
        pk: ', '.join(sorted(subs))
        for pk, subs in staff_subarea_map_raw.items()
    }

    return {
        'laporan_list'         : laporan_list,
        'staff_list'           : staff_list,
        'absensi_list'         : absensi_list,
        'item_list'            : item_list,
        'supervisor_ids'       : supervisor_ids,
        'staff_ids'            : staff_ids,
        'jadwal_kalender'      : jadwal_kalender,
        'absensi_kalender'     : absensi_kalender,
        'laporan_dengan_jadwal': laporan_dengan_jadwal,
        'days_in_month'        : days_in_month,
        'days_range'           : list(range(1, days_in_month + 1)),
        # info_hari untuk highlight header kolom di template
        # Format: {1: {'is_libur': True, 'is_minggu': True, 'is_nasional': False}, ...}
        'info_hari'            : info_hari,
        # staff_subarea_map untuk section I Data Karyawan
        # Format: {staff_pk: 'Sub Area A, Sub Area B'}
        'staff_subarea_map'    : staff_subarea_map,
    }   


# ──────────────────────────────────────────────────
# VIEW UTAMA
# ──────────────────────────────────────────────────

@login_required
def generate_laporan_bulanan(request, perusahaan_id, tahun, bulan, jenis_jasa_id, format):
    try:
        user       = request.user
        perusahaan = get_object_or_404(Perusahaan, pk=perusahaan_id)
        jenis_jasa = get_object_or_404(JenisJasa, pk=jenis_jasa_id)

        if not (user.is_admin or user.is_kepala_supervisor):
            is_assigned = SupervisorPerusahaan.objects.filter(
                supervisor=user,
                perusahaan=perusahaan,
                jenis_jasa=jenis_jasa,
                is_active=True,
            ).exists()
            if not is_assigned:
                return HttpResponse("Tidak punya akses.", status=403)

        data = get_data_laporan_bulanan(perusahaan_id, bulan, tahun, jenis_jasa_id)
        for item in data['item_list']:
            item.foto_on_progress_path = (
                'file://' + item.foto_on_progress.path
                if item.foto_on_progress else None
            )
            item.foto_after_path = (
                'file://' + item.foto_after.path
                if item.foto_after else None
            )

        nama_bulan_list = [
            '', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
            'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
        ]
        nama_bulan = nama_bulan_list[int(bulan)]

        nama_perusahaan_safe = "".join(
            c for c in perusahaan.nama_perusahaan if c.isalnum() or c in (' ', '-', '_')
        ).strip().replace(' ', '_')

        context = {
            **data,
            'perusahaan': perusahaan,
            'jenis_jasa': jenis_jasa,
            'bulan'     : bulan,
            'tahun'     : tahun,
            'nama_bulan': nama_bulan,
            'nama_file' : f"Laporan_{nama_perusahaan_safe}_{nama_bulan}_{tahun}",
        }

        if format == 'pdf':
            return _generate_pdf(request, context)
        elif format == 'word':
            return _generate_word(context)
        else:
            return HttpResponse("Format tidak dikenal. Gunakan 'pdf' atau 'word'.", status=400)

    except Exception as e:
        return HttpResponse(f"Gagal generate laporan: {e}", status=500)


# ──────────────────────────────────────────────────
# PDF & WORD GENERATOR
# ──────────────────────────────────────────────────

def _fetch_resources(uri, rel):
    if uri.startswith(settings.MEDIA_URL):
        return os.path.join(
            settings.MEDIA_ROOT,
            uri.replace(settings.MEDIA_URL, '').lstrip('/')
        )
    if uri.startswith(settings.STATIC_URL):
        static_root = getattr(settings, 'STATIC_ROOT', None) or ''
        return os.path.join(
            static_root,
            uri.replace(settings.STATIC_URL, '').lstrip('/')
        )
    return uri


def _generate_pdf(request, context):
    try:
        from weasyprint import HTML
        import os
        from django.conf import settings

        html_string = render_to_string('laporan_bulanan_pdf.html', context, request=request)
        
        # Ganti ini — baca file dari disk langsung, bukan via HTTP
        base_url = 'file://' + os.path.abspath(settings.BASE_DIR) + '/'
        pdf_bytes = HTML(string=html_string, base_url=base_url).write_pdf()

        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{context["nama_file"]}.pdf"'
        return response

    except Exception as e:
        return HttpResponse(f"Error PDF: {e}", status=500)

def _generate_word(context):
    try:
        from docx.shared import Mm
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Mm(20)
            section.bottom_margin = Mm(20)
            section.left_margin   = Mm(18)
            section.right_margin  = Mm(18)

        _word_cover(doc, context)
        _word_data_karyawan(doc, context)
        _word_struktur_organisasi(doc, context)
        _word_jadwal(doc, context)
        _word_absensi(doc, context)
        _word_program_pelaksanaan(doc, context)
        _word_foto_progres(doc, context)
        _word_penutup(doc, context)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="{context["nama_file"]}.docx"'
        return response

    except Exception as e:
        import traceback  # ← tambah ini
        return HttpResponse(f"Error Word:\n{traceback.format_exc()}", status=500)


# ──────────────────────────────────────────────────
# WORD HELPERS
# ──────────────────────────────────────────────────

def _add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _add_table_header(table, headers):
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(7)

        tc   = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1a3a5c')
        tcPr.append(shd)

        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _word_cover(doc, ctx):
    from docx.shared import Pt

    doc.add_paragraph()
    t = doc.add_paragraph('LAPORAN BULANAN')
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(18)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ctx['perusahaan'].nama_perusahaan + '\n')
    r.bold = True
    r.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Jasa: {ctx['jenis_jasa'].nama_jasa}\n")
    p2.add_run(f"Periode: {ctx['nama_bulan']} {ctx['tahun']}\n")
    p2.add_run(f"\n{ctx['perusahaan'].alamat}")

    doc.add_page_break()


def _word_data_karyawan(doc, ctx):
    _add_section_heading(doc, 'I. DATA KARYAWAN')

    # Build mapping staff_id → sub_area dari item_list
    # Gunakan .get() dengan default [] agar tidak error kalau key tidak ada
    staff_subarea_map = {}
    try:
        for item in ctx.get('item_list', []):
            if item is None:
                continue
            sub_area_nama = (
                item.sub_area.nama_sub_area
                if item.sub_area else '-'
            )
            # item.staff adalah ManyToMany — bisa kosong tapi tidak None
            for staff in item.staff.all():
                if staff is None:
                    continue
                if staff.pk not in staff_subarea_map:
                    staff_subarea_map[staff.pk] = set()
                staff_subarea_map[staff.pk].add(sub_area_nama)
    except Exception:
        # Kalau gagal build map, kolom sub_area akan tampil '-'
        staff_subarea_map = {}

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    _add_table_header(table, ['No', 'Nama', 'NIK', 'Sub Area', 'Status'])

    for i, rel in enumerate(ctx.get('staff_list', []), 1):
        if rel is None:
            continue

        sub_areas    = staff_subarea_map.get(rel.staff.pk, set())
        sub_area_text = ', '.join(sorted(sub_areas)) if sub_areas else '-'

        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = rel.staff.nama_lengkap or rel.staff.username or '-'
        row[2].text = rel.staff.nik or '-'
        row[3].text = sub_area_text
        row[4].text = 'Aktif'

    doc.add_page_break()


def _word_struktur_organisasi(doc, ctx):
    _add_section_heading(doc, 'II. STRUKTUR ORGANISASI')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    _add_table_header(table, ['Supervisor', 'Jenis Jasa', 'Jumlah Staff', 'Status'])

    for laporan in ctx['laporan_list']:
        row = table.add_row().cells
        row[0].text = laporan.supervisor.nama_lengkap or laporan.supervisor.username
        row[1].text = laporan.jenis_jasa.nama_jasa
        row[2].text = str(laporan.supervisor.staff_dibawahnya.count())
        row[3].text = 'Aktif'

    doc.add_page_break()


from docx.shared import Pt, Mm, RGBColor, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _set_col_width(cell, twips_val):
    """Set lebar kolom secara paksa via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(twips_val)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)


def _disable_autofit(table):
    """Matikan autofit agar lebar kolom yang di-set tidak di-override Word."""
    tbl    = table._tbl
    tblPr  = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Hapus tblW lama
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    # Set layout fixed
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)


def _set_row_height(row, height_twips):
    """Set tinggi row agar konsisten."""
    tr    = row._tr
    trPr  = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),  str(int(height_twips)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)


def _header_cell(cell, text, fill='1a3a5c', font_size=5.5):
    """Style cell header: background biru, teks putih bold."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _body_cell(cell, text, font_size=5.5, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False):
    """Style cell body."""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text) if text is not None else '')
    run.font.size = Pt(font_size)
    run.bold = bold
    para.alignment = align


# ── 1 mm ≈ 56.7 twips ──────────────────────────────────────────────────────
MM = 56.7


def _word_jadwal(doc, ctx):
    doc.add_heading('III. JADWAL DAN PLOTING KERJA', level=1)

    days_range      = ctx['days_range']
    jadwal_kalender = ctx['jadwal_kalender']
    n_days          = len(days_range)

    # Usable width A4 portrait dengan margin 18mm kiri-kanan = 210 - 36 = 174mm
    # No=7, Item=30, Task=28, Staff=28 → fixed=93mm, sisa untuk hari
    USABLE  = 174
    FIXED   = 7 + 30 + 28 + 28   # = 93mm
    day_w   = max((USABLE - FIXED) / n_days, 4.5)  # mm per kolom hari

    col_widths_mm = [7, 30, 28, 28] + [day_w] * n_days

    for sub_area_name, items in jadwal_kalender.items():
        p = doc.add_paragraph(sub_area_name)
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(8)

        cols  = 4 + n_days
        table = doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
        _disable_autofit(table)

        # Header
        headers = ['No', 'Item', 'Task', 'Staff'] + [str(d) for d in days_range]
        hdr_row = table.rows[0]
        _set_row_height(hdr_row, 200)
        for i, (h, w) in enumerate(zip(headers, col_widths_mm)):
            cell = hdr_row.cells[i]
            _set_col_width(cell, w * MM)
            _header_cell(cell, h, font_size=5.5)

        # Body
        for idx, item in enumerate(items, 1):
            row = table.add_row()
            _set_row_height(row, 200)
            cells = row.cells

            _set_col_width(cells[0], col_widths_mm[0] * MM)
            _body_cell(cells[0], str(idx), font_size=5.5)

            _set_col_width(cells[1], col_widths_mm[1] * MM)
            _body_cell(cells[1], item['nama_item'], font_size=5.5, align=WD_ALIGN_PARAGRAPH.LEFT)

            _set_col_width(cells[2], col_widths_mm[2] * MM)
            _body_cell(cells[2], item['task'], font_size=5.5, align=WD_ALIGN_PARAGRAPH.LEFT)

            _set_col_width(cells[3], col_widths_mm[3] * MM)
            _body_cell(cells[3], item['staff'], font_size=5.5, align=WD_ALIGN_PARAGRAPH.LEFT)

            for k, day in enumerate(days_range):
                ci  = 4 + k
                val = 'A' if day in item['tanggal_aktif'] else ''
                _set_col_width(cells[ci], day_w * MM)
                _body_cell(cells[ci], val, font_size=5)

        doc.add_paragraph()


def _word_absensi(doc, ctx):
    doc.add_heading('IV. ABSENSI', level=1)

    days_range       = ctx['days_range']
    absensi_kalender = ctx['absensi_kalender']
    n_days           = len(days_range)

    # No=6, Nama=30, NIK=18, hari×X, E/A/I/L/DC=6 (×5=30), Ket=18
    # 174 - 6 - 30 - 18 - 30 - 18 = 72mm untuk hari
    USABLE   = 174
    FIXED    = 6 + 30 + 18 + 30 + 18   # = 102mm
    day_w    = max((USABLE - FIXED) / n_days, 4.0)  # mm per kolom hari

    sum_w  = 6.0   # lebar kolom summary (E A I L DC)
    ket_w  = 18.0

    col_widths_mm = [6, 30, 18] + [day_w] * n_days + [sum_w]*5 + [ket_w]

    cols  = 3 + n_days + 5 + 1
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    _disable_autofit(table)

    # Header
    headers = ['No', 'Nama', 'NIK'] + [str(d) for d in days_range] + ['E', 'A', 'I', 'L', 'DC', 'Ket']
    hdr_row = table.rows[0]
    _set_row_height(hdr_row, 200)
    for i, (h, w) in enumerate(zip(headers, col_widths_mm)):
        cell = hdr_row.cells[i]
        _set_col_width(cell, w * MM)
        _header_cell(cell, h, font_size=5.5)

    # Body
    for i, (staff_id, data) in enumerate(absensi_kalender.items(), 1):
        row = table.add_row()
        _set_row_height(row, 180)
        cells = row.cells

        # No
        _set_col_width(cells[0], col_widths_mm[0] * MM)
        _body_cell(cells[0], str(i), font_size=5.5)

        # Nama
        _set_col_width(cells[1], col_widths_mm[1] * MM)
        _body_cell(cells[1], data['nama'], font_size=5.5, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)

        # NIK
        _set_col_width(cells[2], col_widths_mm[2] * MM)
        _body_cell(cells[2], data['nik'] or '-', font_size=5.5)

        # Kolom hari
        for k in range(n_days):
            ci        = 3 + k
            hari_info = data['hari_list'][k]
            status    = hari_info['status'] if isinstance(hari_info, dict) else (hari_info or '')
            _set_col_width(cells[ci], day_w * MM)
            _body_cell(cells[ci], status, font_size=5)

        # Summary E A I L DC Ket
        base    = 3 + n_days
        summary = [
            data.get('jumlah_hadir',  0),
            data.get('jumlah_alpa',   0),
            data.get('jumlah_izin',   0),
            data.get('jumlah_cuti',   0),
            data.get('jumlah_dokter', 0),
            '',
        ]
        sum_widths = [sum_w]*5 + [ket_w]
        for m, (val, w) in enumerate(zip(summary, sum_widths)):
            _set_col_width(cells[base + m], w * MM)
            _body_cell(cells[base + m], str(val), font_size=5.5)

    doc.add_paragraph()
    ket = doc.add_paragraph()
    kr  = ket.add_run('Keterangan: ')
    kr.bold = True
    kr.font.size = Pt(7)
    kt = ket.add_run('P=Hadir  L=Cuti  I=Izin  A=Alpa  DC=Surat Dokter  LB=Libur')
    kt.font.size = Pt(7)

def _word_program_pelaksanaan(doc, ctx):
    _add_section_heading(doc, 'V. PROGRAM DAN PELAKSANAAN PEKERJAAN')

    days_range            = ctx['days_range']
    laporan_dengan_jadwal = ctx['laporan_dengan_jadwal']

    for entry in laporan_dengan_jadwal:
        p = doc.add_paragraph(f"{entry['area_nama']} — Supervisor: {entry['supervisor']}")
        if p.runs:
            p.runs[0].bold = True

        cols  = 4 + len(days_range)
        table = doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
        _add_table_header(
            table,
            ['Object', 'Standar', 'Job', 'Frek'] + [str(d) for d in days_range]
        )

        for item in entry['items']:
            row = table.add_row().cells
            row[0].text = item['nama_item']
            row[1].text = item['sub_area']
            row[2].text = item['task']
            row[3].text = item['frek']
            for j, mark in enumerate(item['hari_list']):
                row[4 + j].text = mark

        doc.add_paragraph()

    doc.add_page_break()


def _word_foto_progres(doc, ctx):
    _add_section_heading(doc, 'VI. FOTO PROGRES')

    current_area = None
    for item in ctx['item_list']:
        if not (item.foto_on_progress or item.foto_after):
            continue

        area_name = item.laporan.area.nama_area if item.laporan.area else 'Tanpa Area'
        if area_name != current_area:
            p = doc.add_paragraph(area_name)
            if p.runs:
                p.runs[0].bold = True
            current_area = area_name

        doc.add_paragraph(
            f"{item.nama_item}"
            + (f" — {item.tanggal.strftime('%d %B %Y')}" if item.tanggal else "")
        )

        if item.foto_on_progress:
            try:
                doc.add_picture(item.foto_on_progress.path, width=Inches(2.8))
                doc.add_paragraph('On Progress')
            except Exception:
                doc.add_paragraph('[Foto On Progress tidak tersedia]')

        if item.foto_after:
            try:
                doc.add_picture(item.foto_after.path, width=Inches(2.8))
                doc.add_paragraph('After')
            except Exception:
                doc.add_paragraph('[Foto After tidak tersedia]')

        doc.add_paragraph()

    doc.add_page_break()


def _word_penutup(doc, ctx):
    from docx.shared import Pt

    _add_section_heading(doc, 'VII. PENUTUP')
    doc.add_paragraph(
        f"Demikian Laporan Bulanan periode bulan {ctx['nama_bulan']} {ctx['tahun']} ini dibuat "
        f"sebagai bentuk pertanggungjawaban pelaksanaan jasa {ctx['jenis_jasa'].nama_jasa} "
        f"di lingkungan {ctx['perusahaan'].nama_perusahaan}. "
        f"Semua kegiatan telah dilaksanakan sesuai dengan jadwal dan standar yang telah ditetapkan."
    )

    doc.add_paragraph()
    p = doc.add_paragraph(f"{ctx['perusahaan'].alamat or ''}, {ctx['nama_bulan']} {ctx['tahun']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    supervisor_name = ''
    for laporan in ctx['laporan_list']:
        supervisor_name = laporan.supervisor.nama_lengkap or laporan.supervisor.username
        break

    table.cell(0, 0).text = 'Dibuat oleh,'
    table.cell(0, 1).text = 'Diketahui oleh,'
    table.cell(0, 2).text = 'Disetujui oleh,'
    table.cell(1, 0).text = f'( {supervisor_name} )'
    table.cell(1, 1).text = '( ________________________ )'
    table.cell(1, 2).text = '( ________________________ )'
    table.cell(2, 0).text = f'SPV {ctx["jenis_jasa"].nama_jasa}'
    table.cell(2, 1).text = 'Kepala Supervisor'
    table.cell(2, 2).text = ctx['perusahaan'].nama_perusahaan